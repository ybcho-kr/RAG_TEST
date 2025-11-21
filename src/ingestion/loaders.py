"""다양한 문서 포맷을 읽어 `Document` 객체로 변환하는 로더 모음."""
import importlib.util
from pathlib import Path
from typing import Iterable, List

from src.models.document import Document


class UnsupportedFormatError(ValueError):
    """지원하지 않는 파일 형식일 때 발생하는 예외."""


def load_txt(path: Path) -> List[Document]:
    """평문 텍스트 파일을 읽어 Document 리스트로 반환한다."""

    text = path.read_text(encoding="utf-8")
    return [Document(content=text, source=str(path))]


def load_pdf(path: Path) -> List[Document]:
    """PDF 파일을 페이지 단위로 파싱한다."""

    if importlib.util.find_spec("pypdf") is None:
        raise ImportError("pypdf 패키지가 필요합니다. 설치 후 다시 실행하세요.")

    from pypdf import PdfReader

    reader = PdfReader(str(path))
    documents: List[Document] = []
    for idx, page in enumerate(reader.pages):
        text = page.extract_text() or ""
        documents.append(Document(content=text, source=str(path), page=idx + 1))
    return documents


def load_docx(path: Path) -> List[Document]:
    """DOCX 파일을 단일 텍스트로 정규화한다."""

    if importlib.util.find_spec("docx") is None:
        raise ImportError("python-docx 패키지가 필요합니다. 설치 후 다시 실행하세요.")

    import docx

    document = docx.Document(str(path))
    paragraphs = [p.text for p in document.paragraphs if p.text]
    text = "\n".join(paragraphs)
    return [Document(content=text, source=str(path))]


def load_html(path: Path) -> List[Document]:
    """HTML 파일에서 가시 텍스트를 추출한다."""

    if importlib.util.find_spec("bs4") is None:
        raise ImportError("beautifulsoup4 패키지가 필요합니다. 설치 후 다시 실행하세요.")

    from bs4 import BeautifulSoup

    html = path.read_text(encoding="utf-8")
    soup = BeautifulSoup(html, "html.parser")
    for tag in soup(["script", "style"]):
        tag.decompose()
    text = "\n".join(chunk.strip() for chunk in soup.get_text("\n").splitlines() if chunk.strip())
    return [Document(content=text, source=str(path))]


LOADERS = {
    ".txt": load_txt,
    ".pdf": load_pdf,
    ".docx": load_docx,
    ".html": load_html,
    ".htm": load_html,
}


def load_documents(paths: Iterable[str]) -> List[Document]:
    """파일 경로들의 이터러블을 받아 문서 리스트를 반환한다."""

    documents: List[Document] = []
    for raw_path in paths:
        path = Path(raw_path)
        if not path.exists():
            raise FileNotFoundError(f"파일을 찾을 수 없습니다: {raw_path}")

        loader = LOADERS.get(path.suffix.lower())
        if not loader:
            raise UnsupportedFormatError(f"지원하지 않는 파일 형식입니다: {path.suffix}")

        documents.extend(loader(path))
    return documents
