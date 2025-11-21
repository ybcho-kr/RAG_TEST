"""DOCX ingestion utilities using python-docx."""
from pathlib import Path

from .document import Document

try:
    import docx
except Exception:  # pragma: no cover - optional dependency
    docx = None  # type: ignore


def load_docx(path: str) -> Document:
    """Load a DOCX file into a :class:`Document`."""
    if docx is None:
        raise ImportError("python-docx is required to load DOCX files")

    file_path = Path(path)
    doc = docx.Document(str(file_path))
    paragraphs = [p.text for p in doc.paragraphs]
    content = "\n".join(filter(None, paragraphs))
    return Document(content=content, source=str(file_path), metadata={"format": "docx"})
